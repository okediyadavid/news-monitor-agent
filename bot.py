"""
Telegram bot command handler for news monitoring agent.
Allows users to interact with the system directly through Telegram.
Uses simple polling with requests library for Python 3.14 compatibility.
"""

import logging
import requests
import time
import hashlib
import os
import re
import threading
from typing import Optional, Dict, Any
from datetime import datetime
from http.server import HTTPServer, BaseHTTPRequestHandler

from database import DatabaseManager
from scraper import WebScraper
from rss import RSSParser
from notifier import TelegramNotifier
from scheduler import NewsScheduler

logger = logging.getLogger(__name__)


class HealthCheckHandler(BaseHTTPRequestHandler):
    """Simple HTTP handler for health checks."""
    
    def do_GET(self):
        self.send_response(200)
        self.send_header('Content-type', 'text/plain')
        self.end_headers()
        self.wfile.write(b'OK')
    
    def log_message(self, format, *args):
        """Suppress default logging."""
        pass


class NewsBot:
    """Telegram bot for news monitoring using simple polling."""
    
    def __init__(self, bot_token: str, db_path: str = "data/news_monitor.db"):
        """
        Initialize the Telegram bot.
        
        Args:
            bot_token: Telegram bot token
            db_path: Path to database file
        """
        self.bot_token = bot_token
        self.api_url = f"https://api.telegram.org/bot{bot_token}"
        self.db = DatabaseManager(db_path)
        self.web_scraper = WebScraper()
        self.rss_parser = RSSParser()
        self.notifier = TelegramNotifier(bot_token, "")  # Chat ID will be set per user
        self.offset = 0
        
        # Conversation state management
        self.user_states = {}
        
        # Pre-configured news sources
        self.pre_initialized_sources = [
            {
                'name': 'TechNext24',
                'url': 'https://technext24.com/',
                'type': 'website',
                'category': 'Technology'
            },
            {
                'name': 'Today Africa',
                'url': 'https://todayafrica.co/',
                'type': 'website',
                'category': 'General'
            },
            {
                'name': 'CNBC Africa',
                'url': 'https://www.cnbcafrica.com/',
                'type': 'website',
                'category': 'Finance'
            },
            {
                'name': 'Cryptopolitan',
                'url': 'https://www.cryptopolitan.com/',
                'type': 'website',
                'category': 'Crypto'
            }
        ]
        
        # Initialize scheduler (default 6 hours)
        self.scheduler = NewsScheduler(check_interval_hours=6)
        self.scheduler.set_check_function(self._scheduled_check)
        
        self._initialize_preconfigured_sources()
    
    def _initialize_preconfigured_sources(self) -> None:
        """Initialize pre-configured news sources for new users."""
        # These will be added when a user registers
        logger.info(f"Pre-configured sources: {len(self.pre_initialized_sources)}")
    
    def get_updates(self) -> list:
        """Get updates from Telegram."""
        try:
            params = {'offset': self.offset, 'timeout': 10}
            response = requests.get(f"{self.api_url}/getUpdates", params=params, timeout=15)
            response.raise_for_status()
            data = response.json()
            
            if data.get('ok'):
                updates = data.get('result', [])
                if updates:
                    self.offset = updates[-1]['update_id'] + 1
                    logger.info(f"Got {len(updates)} updates, new offset: {self.offset}")
                return updates
            return []
        except Exception as e:
            logger.error(f"Error getting updates: {e}")
            return []
    
    def send_message(self, chat_id: int, text: str, parse_mode: str = None, max_retries: int = 3) -> bool:
        """Send a message to a chat with retry logic."""
        from tenacity import retry, stop_after_attempt, wait_exponential
        
        @retry(stop=stop_after_attempt(max_retries), wait=wait_exponential(multiplier=1, min=2, max=10))
        def _send():
            data = {'chat_id': chat_id, 'text': text}
            if parse_mode:
                data['parse_mode'] = parse_mode
            
            logger.info(f"Sending message to {chat_id}: {text[:50]}...")
            response = requests.post(f"{self.api_url}/sendMessage", json=data, timeout=10)
            response.raise_for_status()
            result = response.json()
            logger.info(f"Message sent: {result.get('ok', False)}")
            return result.get('ok', False)
        
        try:
            return _send()
        except Exception as e:
            logger.error(f"Error sending message after retries: {e}")
            return False
    
    def send_document(self, chat_id: int, filepath: str, filename: str, max_retries: int = 3) -> bool:
        """Send a document to a chat with retry logic."""
        from tenacity import retry, stop_after_attempt, wait_exponential
        
        @retry(stop=stop_after_attempt(max_retries), wait=wait_exponential(multiplier=1, min=2, max=10))
        def _send():
            with open(filepath, 'rb') as f:
                files = {'document': (filename, f)}
                data = {'chat_id': chat_id}
                
                logger.info(f"Sending document to {chat_id}: {filename}")
                response = requests.post(f"{self.api_url}/sendDocument", data=data, files=files, timeout=30)
                response.raise_for_status()
                result = response.json()
                logger.info(f"Document sent: {result.get('ok', False)}")
                return result.get('ok', False)
        
        try:
            return _send()
        except Exception as e:
            logger.error(f"Error sending document after retries: {e}")
            return False
    
    def handle_command(self, chat_id: int, command: str, args: list) -> None:
        """Handle a command from a user."""
        logger.info(f"Handling command: {command} from chat {chat_id}")
        
        if command == 'start':
            self.start_command(chat_id)
        elif command == 'help':
            self.help_command(chat_id)
        elif command == 'register':
            self.register_command(chat_id, args)
        elif command == 'addsource':
            self.add_source_command(chat_id, args)
        elif command == 'mysources':
            self.my_sources_command(chat_id)
        elif command == 'removesource':
            self.remove_source_command(chat_id, args)
        elif command == 'setinterval':
            self.set_interval_command(chat_id, args)
        elif command == 'stats':
            self.stats_command(chat_id)
        elif command == 'checknow':
            self.check_now_command(chat_id)
        elif command == 'checktoday':
            self.check_today_command(chat_id)
        elif command == 'checkyesterday':
            self.check_yesterday_command(chat_id)
        elif command == 'check12h':
            self.check_12h_command(chat_id)
        elif command == 'check6h':
            self.check_6h_command(chat_id)
        elif command == 'check1h':
            self.check_1h_command(chat_id)
        elif command == 'addkeyword':
            self.add_keyword_command(chat_id, args)
        elif command == 'removekeyword':
            self.remove_keyword_command(chat_id, args)
        elif command == 'mykeywords':
            self.my_keywords_command(chat_id)
        elif command == 'bookmark':
            self.bookmark_command(chat_id, args)
        elif command == 'unbookmark':
            self.unbookmark_command(chat_id, args)
        elif command == 'mybookmarks':
            self.my_bookmarks_command(chat_id)
        elif command == 'search':
            self.search_command(chat_id, args)
        elif command == 'searchfrom':
            self.search_from_command(chat_id, args)
        elif command == 'startscheduler':
            self.start_scheduler_command(chat_id)
        elif command == 'stopscheduler':
            self.stop_scheduler_command(chat_id)
        elif command == 'schedulerstatus':
            self.scheduler_status_command(chat_id)
        elif command == 'category':
            self.category_command(chat_id, args)
        elif command == 'setcategory':
            self.set_category_command(chat_id, args)
        elif command == 'sourcestats':
            self.source_stats_command(chat_id)
        elif command == 'reliablesources':
            self.reliable_sources_command(chat_id)
        elif command == 'myanalytics':
            self.my_analytics_command(chat_id)
        elif command == 'exporttoday':
            self.export_today_command(chat_id, args)
        elif command == 'exportlast7days':
            self.export_last_7days_command(chat_id, args)
        elif command == 'trending':
            self.trending_command(chat_id)
        elif command == 'summarize':
            self.summarize_command(chat_id, args)
        elif command == 'cancel':
            self.cancel_command(chat_id)
        elif command == 'interest':
            self.interest_command(chat_id, args)
        elif command == 'myinterests':
            self.my_interests_command(chat_id, args)
        elif command == 'uninterest':
            self.uninterest_command(chat_id, args)
        elif command == 'exportdaily':
            self.export_daily_command(chat_id, args)
        else:
            self.send_message(chat_id, "Unknown command. Use /help for available commands.")
    
    def handle_message(self, chat_id: int, text: str) -> None:
        """Handle non-command messages (natural language processing)."""
        logger.info(f"Handling message from {chat_id}: {text}")
        
        # Check if user is in a conversation state
        if chat_id in self.user_states:
            state = self.user_states[chat_id]
            
            if state['state'] == 'awaiting_name':
                self._handle_registration_name(chat_id, text)
            elif state['state'] == 'awaiting_source_url':
                self._handle_source_url(chat_id, text)
            elif state['state'] == 'awaiting_source_name':
                self._handle_source_name(chat_id, text)
            elif state['state'] == 'awaiting_source_category':
                self._handle_source_category(chat_id, text)
            elif state['state'] == 'awaiting_interval':
                self._handle_interval(chat_id, text)
        else:
            # Natural language processing for non-command messages
            self._handle_natural_language(chat_id, text)
    
    def _handle_natural_language(self, chat_id: int, text: str) -> None:
        """Process natural language requests without commands with enhanced pattern matching."""
        text_lower = text.lower().strip()
        
        # Check if user is registered
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            if any(word in text_lower for word in ['register', 'sign up', 'join', 'create account']):
                self.register_command(chat_id)
            else:
                self.send_message(chat_id, "👋 Welcome! Please register first by typing 'register' or use /register")
            return
        
        # Enhanced natural language intent matching with regex patterns
        import re
        
        # Check news intents - more sophisticated patterns
        if re.search(r'(check|show|get|display|list|what are|tell me about).*(news|articles|headlines|updates|stories)', text_lower):
            self.check_now_command(chat_id)
            return
        
        if re.search(r'(today|today\'s|current|latest|recent).*(news|articles|headlines)', text_lower):
            self.check_today_command(chat_id)
            return
        
        if re.search(r'(yesterday|yesterday\'s|past).*(news|articles)', text_lower):
            self.check_yesterday_command(chat_id)
            return
        
        if re.search(r'(last|past).*(\d+)\s*(hours|hr|h)', text_lower):
            match = re.search(r'(\d+)\s*(hours|hr|h)', text_lower)
            if match:
                hours = int(match.group(1))
                if hours <= 1:
                    self.check_1h_command(chat_id)
                elif hours <= 6:
                    self.check_6h_command(chat_id)
                elif hours <= 12:
                    self.check_12h_command(chat_id)
                else:
                    self.check_today_command(chat_id)
            return
        
        # Source management intents
        if re.search(r'(add|create|new).*(source|website|rss|feed|url)', text_lower):
            self.add_source_command(chat_id, [])
            return
        
        if re.search(r'(my|list|show|display).*(sources|websites|feeds)', text_lower):
            self.my_sources_command(chat_id)
            return
        
        if re.search(r'(remove|delete).*(source|website)', text_lower):
            self.send_message(chat_id, "To remove a source, please use /removesource with the source ID")
            return
        
        # Search intents - extract query
        if re.search(r'(search|find|look for|find|search for|look up).*(about|for|on)', text_lower):
            # Extract the search query
            query = re.sub(r'(search|find|look for|find|search for|look up|about|for|on)', '', text_lower).strip()
            if query:
                self.search_command(chat_id, query.split())
            else:
                self.send_message(chat_id, "What would you like to search for?")
            return
        
        # Bookmark intents
        if re.search(r'(bookmark|save).*(article)', text_lower):
            self.send_message(chat_id, "To bookmark an article, please use /bookmark with the article ID")
            return
        
        if re.search(r'(my|show|list).*(bookmarks|saved|saved articles)', text_lower):
            self.my_bookmarks_command(chat_id)
            return
        
        # Keyword intents
        if re.search(r'(add|track follow).*(keyword|topic|interest)', text_lower):
            # Extract keyword
            match = re.search(r'(add|track|follow).*(keyword|topic|interest)\s+(.+)', text_lower)
            if match:
                keyword = match.group(3).strip()
                self.add_keyword_command(chat_id, [keyword])
            else:
                self.send_message(chat_id, "What keyword would you like to add?")
            return
        
        if re.search(r'(my|show|list).*(keywords|topics|interests)', text_lower):
            self.my_keywords_command(chat_id)
            return
        
        # Analytics and stats
        if re.search(r'(stats|statistics|analytics|my stats|performance)', text_lower):
            self.stats_command(chat_id)
            return
        
        if re.search(r'(source|website).*(stats|statistics|performance|reliable)', text_lower):
            self.source_stats_command(chat_id)
            return
        
        # Categories
        if re.search(r'(category|categories|categorize)', text_lower):
            self.category_command(chat_id, [])
            return
        
        # Export
        if re.search(r'(export|download).*(articles|news)', text_lower):
            self.send_message(chat_id, "To export articles, use /exporttoday or /exportlast7days")
            return
        
        # Trending
        if re.search(r'(trending|popular|hot|what\'s trending|what\'s popular)', text_lower):
            self.trending_command(chat_id)
            return
        
        # Help
        if re.search(r'(help|commands|what can you do|how to use|instructions)', text_lower):
            self.help_command(chat_id)
            return
        
        # Question/Answer pattern - try AI if available
        if re.search(r'(what|how|why|when|where|who|tell me|explain|describe)', text_lower):
            self._handle_ai_question(chat_id, text)
            return
        
        # Fallback: try to extract search query from longer text
        if len(text_lower.split()) > 2:
            self.search_command(chat_id, text.split())
        else:
            self.send_message(chat_id, "🤔 I'm not sure what you mean. Try asking me to 'show news', 'search for something', 'add source', or type /help for all commands.")
    
    def _handle_ai_question(self, chat_id: int, question: str) -> None:
        """Handle AI-powered questions about news and articles."""
        try:
            import os
            openai_api_key = os.getenv('OPENAI_API_KEY')
            
            if not openai_api_key:
                # Fallback to search if no AI available
                self.send_message(chat_id, "🤔 I'd love to answer that with AI, but no AI API key is configured. Let me search for related articles instead...")
                self.search_command(chat_id, question.split())
                return
            
            # Get recent articles for context
            user = self._get_user_by_chat_id(chat_id)
            if not user:
                self.send_message(chat_id, "❌ Please register first")
                return
            
            with self.db._get_connection() as conn:
                cursor = conn.cursor()
                cursor.execute("""
                    SELECT a.*, s.name as source_name 
                    FROM articles a 
                    JOIN sources s ON a.source_id = s.id 
                    WHERE a.user_id = ?
                    ORDER BY a.created_at DESC
                    LIMIT 10
                """, (user['id'],))
                
                articles = [dict(row) for row in cursor.fetchall()]
            
            if not articles:
                self.send_message(chat_id, "🤔 I don't have any articles to analyze yet. Try checking for news first with 'show news'")
                return
            
            # Prepare context for AI
            context = "Here are recent news articles:\n\n"
            for i, article in enumerate(articles[:5], 1):
                context += f"{i}. {article['title']}\n"
                context += f"   Source: {article['source_name']}\n"
                if article.get('summary'):
                    context += f"   Summary: {article['summary']}\n"
                context += "\n"
            
            context += f"\nUser question: {question}"
            
            # Call OpenAI API
            import openai
            client = openai.OpenAI(api_key=openai_api_key)
            
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a helpful news assistant. Answer questions about the provided news articles. Be concise and helpful."},
                    {"role": "user", "content": context}
                ],
                max_tokens=500,
                temperature=0.7
            )
            
            answer = response.choices[0].message.content
            self.send_message(chat_id, f"🤖 AI Answer:\n\n{answer}")
            
        except Exception as e:
            logger.error(f"Error in AI question handling: {e}")
            self.send_message(chat_id, "🤔 I had trouble processing that with AI. Let me search for related articles instead...")
            self.search_command(chat_id, question.split())
    
    def _ai_summarize_articles(self, chat_id: int, article_ids: list) -> None:
        """Use AI to intelligently summarize multiple articles."""
        try:
            import os
            openai_api_key = os.getenv('OPENAI_API_KEY')
            
            if not openai_api_key:
                self.send_message(chat_id, "🤔 AI summarization requires an OpenAI API key. Add OPENAI_API_KEY to your .env file.")
                return
            
            user = self._get_user_by_chat_id(chat_id)
            if not user:
                self.send_message(chat_id, "❌ Please register first")
                return
            
            # Get articles
            with self.db._get_connection() as conn:
                cursor = conn.cursor()
                placeholders = ','.join(['?' for _ in article_ids])
                cursor.execute(f"""
                    SELECT a.*, s.name as source_name 
                    FROM articles a 
                    JOIN sources s ON a.source_id = s.id 
                    WHERE a.id IN ({placeholders}) AND a.user_id = ?
                """, article_ids + [user['id']])
                
                articles = [dict(row) for row in cursor.fetchall()]
            
            if not articles:
                self.send_message(chat_id, "❌ No articles found to summarize")
                return
            
            # Prepare context for AI
            context = "Please provide a concise summary of these news articles, highlighting key themes and connections:\n\n"
            for i, article in enumerate(articles, 1):
                context += f"{i}. {article['title']}\n"
                context += f"   Source: {article['source_name']}\n"
                if article.get('summary'):
                    context += f"   Summary: {article['summary']}\n"
                context += "\n"
            
            # Call OpenAI API
            import openai
            client = openai.OpenAI(api_key=openai_api_key)
            
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a news summarization expert. Provide concise, insightful summaries that highlight key themes and connections between articles."},
                    {"role": "user", "content": context}
                ],
                max_tokens=600,
                temperature=0.5
            )
            
            summary = response.choices[0].message.content
            self.send_message(chat_id, f"📝 AI Summary:\n\n{summary}")
            
        except Exception as e:
            logger.error(f"Error in AI summarization: {e}")
            self.send_message(chat_id, "🤔 I had trouble generating the AI summary.")
    
    def _execute_intent(self, chat_id: int, intent: str, original_text: str) -> None:
        """Execute the matched natural language intent."""
        text_lower = original_text.lower()
        
        if intent == 'check news':
            self.check_now_command(chat_id)
        elif intent == 'check today':
            self.check_today_command(chat_id)
        elif intent == 'check yesterday':
            self.check_yesterday_command(chat_id)
        elif intent == 'add source':
            self.add_source_command(chat_id, [])
        elif intent == 'my sources':
            self.my_sources_command(chat_id)
        elif intent == 'remove source':
            self.send_message(chat_id, "To remove a source, please use /removesource with the source ID")
        elif intent == 'search':
            # Extract search query
            words = original_text.split()
            if len(words) > 1:
                self.search_command(chat_id, words[1:])
            else:
                self.send_message(chat_id, "What would you like to search for?")
        elif intent == 'bookmark':
            self.send_message(chat_id, "To bookmark an article, please use /bookmark with the article ID")
        elif intent == 'my bookmarks':
            self.my_bookmarks_command(chat_id)
        elif intent == 'add keyword':
            # Extract keyword
            words = original_text.split()
            if len(words) > 2:
                keyword = ' '.join(words[2:])
                self.add_keyword_command(chat_id, [keyword])
            else:
                self.send_message(chat_id, "What keyword would you like to add?")
        elif intent == 'my keywords':
            self.my_keywords_command(chat_id)
        elif intent == 'stats':
            self.stats_command(chat_id)
        elif intent == 'source stats':
            self.source_stats_command(chat_id)
        elif intent == 'category':
            self.category_command(chat_id, [])
        elif intent == 'export':
            self.send_message(chat_id, "To export articles, use /exporttoday or /exportlast7days")
        elif intent == 'trending':
            self.trending_command(chat_id)
        elif intent == 'help':
            self.help_command(chat_id)
    
    def start_command(self, chat_id: int) -> None:
        """Handle /start command."""
        user = self._get_user_by_chat_id(chat_id)
        
        if user:
            message = f"""👋 Welcome back, {user['email']}!

I'm your News Monitor Bot. Here's what you can do:

📰 Manage News Sources
/addsource - Add a new news source
/mysources - View your news sources
/removesource - Remove a news source

⚙️ Settings
/setinterval - Change check interval (hours)

📊 Information
/stats - View your statistics
/checknow - Check for new articles now
/help - Show all commands

Your current check interval: 6 hours"""
        else:
            message = """👋 Welcome to News Monitor Bot!

I'll help you monitor news sources and send you updates.

To get started, register with:
/register your@email.com

After registration, you can:
• Add news sources to monitor
• Set your preferred check interval
• Receive notifications for new articles

Type /help for all available commands."""
        
        self.send_message(chat_id, message)
    
    def help_command(self, chat_id: int) -> None:
        """Handle /help command."""
        logger.info(f"Executing help_command for chat {chat_id}")
        message = """📚 News Monitor Bot Commands

User Management
/register - Register with your name (interactive)
/start - Start the bot / show welcome

News Sources
/addsource - Add a new news source (interactive with URL analysis)
/mysources - View your news sources
/removesource source_id - Remove a news source

Article Checking
/checknow - Check for new articles immediately
/checktoday - Check all articles from today
/checkyesterday - Check articles from yesterday
/check12h - Check articles from last 12 hours
/check6h - Check articles from last 6 hours
/check1h - Check articles from last 1 hour

Keyword Filtering
/addkeyword keyword - Add keyword for filtering
/removekeyword keyword - Remove keyword
/mykeywords - View your keywords

Bookmarks
/bookmark article_id - Save article for later
/unbookmark article_id - Remove bookmark
/mybookmarks - View saved articles

Search
/search query - Search articles by title/summary
/searchfrom source_name query - Search specific source

Scheduler
/startscheduler - Start automatic article checking
/stopscheduler - Stop automatic article checking
/schedulerstatus - View scheduler status

Categories
/category - View articles by category
/setcategory source_id category - Set category for source

Performance
/sourcestats - View source performance metrics
/reliablesources - Show most reliable sources

Analytics
/myanalytics - View your analytics dashboard

Export
/exporttoday json/csv - Export today's articles
/exportlast7days json/csv - Export last 7 days articles

Trending
/trending - View trending topics

AI Features (requires OPENAI_API_KEY)
/summarize - AI-powered article summarization
"what is..." - Ask questions about news (AI-powered)

Interest Catalog
/interest article_id - Mark an article as interesting (with AI summary)
/myinterests [date] - View your daily interests (default: today)
/uninterest article_id - Remove article from interests
/exportdaily [date] - Generate Word document with daily interests

Other
/cancel - Cancel current operation
/help - Show this help message

💡 Interactive Mode:
Commands like /register, /addsource, and /setinterval now work interactively!
Just type the command and follow the prompts - no need to type everything at once."""
        
        logger.info(f"Calling send_message for help_command")
        result = self.send_message(chat_id, message)
        logger.info(f"send_message returned: {result}")
    
    def register_command(self, chat_id: int, args: list) -> None:
        """Handle /register command - interactive mode."""
        # Check if already registered
        existing_user = self._get_user_by_chat_id(chat_id)
        if existing_user:
            self.send_message(chat_id, f"❌ You're already registered as {existing_user['name']}")
            return
        
        # Start interactive registration
        self.user_states[chat_id] = {'state': 'awaiting_name'}
        self.send_message(chat_id, "👤 Please enter your name to register:")
    
    def _handle_registration_name(self, chat_id: int, text: str) -> None:
        """Handle name input during registration."""
        name = text.strip()
        
        # Validate name
        if len(name) < 2:
            self.send_message(chat_id, "❌ Name must be at least 2 characters. Please try again:")
            return
        
        # Check if name already exists
        if self._get_user_by_name(name):
            self.send_message(chat_id, "❌ This name is already registered. Please use a different name:")
            return
        
        # Create user
        try:
            user_id = self._create_user(name, chat_id)
            # Clear state
            del self.user_states[chat_id]
            
            # Add pre-configured sources
            self._add_preconfigured_sources(user_id)
            
            self.send_message(chat_id, f"✅ Successfully registered!\n\nName: {name}\n\nI've added these pre-configured news sources for you:\n\n" + 
                              "\n".join([f"• {s['name']}: {s['url']}" for s in self.pre_initialized_sources]) +
                              "\n\nYou can add more sources with /addsource or view your sources with /mysources")
        except Exception as e:
            logger.error(f"Error registering user: {e}")
            self.send_message(chat_id, "❌ Registration failed. Please try again.")
            del self.user_states[chat_id]
    
    def _add_preconfigured_sources(self, user_id: int) -> None:
        """Add pre-configured sources for a new user."""
        for source in self.pre_initialized_sources:
            try:
                self.db.add_source(
                    name=source['name'],
                    url=source['url'],
                    source_type='website',
                    category=source['category'],
                    user_id=user_id
                )
                logger.info(f"Added pre-configured source {source['name']} for user {user_id}")
            except Exception as e:
                logger.error(f"Error adding pre-configured source {source['name']}: {e}")
    
    def add_source_command(self, chat_id: int, args: list) -> None:
        """Handle /addsource command - interactive mode."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register")
            return
        
        # Start interactive source addition
        self.user_states[chat_id] = {'state': 'awaiting_source_url', 'user_id': user['id']}
        self.send_message(chat_id, "📰 Please enter the news source URL:\n\nI'll analyze it and help you set it up.")
    
    def _handle_source_url(self, chat_id: int, text: str) -> None:
        """Handle URL input during source addition."""
        url = text.strip()
        
        # Validate URL
        if not url.startswith('http://') and not url.startswith('https://'):
            self.send_message(chat_id, "❌ Invalid URL. Please include http:// or https://")
            return
        
        # Analyze URL
        analysis = self._analyze_url(url)
        
        # Store URL and move to name step
        self.user_states[chat_id]['url'] = url
        self.user_states[chat_id]['analysis'] = analysis
        self.user_states[chat_id]['state'] = 'awaiting_source_name'
        
        message = f"✅ URL validated!\n\n📊 Analysis:\n{analysis}\n\nPlease enter a name for this source (or press Enter to use suggested name):"
        self.send_message(chat_id, message)
    
    def _analyze_url(self, url: str) -> str:
        """Analyze a news source URL and provide information."""
        analysis = []
        
        try:
            from urllib.parse import urlparse
            parsed = urlparse(url)
            domain = parsed.netloc.replace('www.', '')
            
            analysis.append(f"• Domain: {domain}")
            analysis.append(f"• Protocol: {parsed.scheme}")
            
            # Try to fetch and analyze
            try:
                response = requests.head(url, timeout=5, allow_redirects=True)
                if response.status_code == 200:
                    analysis.append(f"• Status: ✅ Accessible")
                    content_type = response.headers.get('content-type', '')
                    if 'html' in content_type:
                        analysis.append(f"• Type: HTML Website")
                    elif 'xml' in content_type or 'rss' in content_type:
                        analysis.append(f"• Type: RSS Feed")
                else:
                    analysis.append(f"• Status: ⚠️ HTTP {response.status_code}")
            except:
                analysis.append(f"• Status: ⚠️ Could not verify accessibility")
            
            # Suggest name
            suggested_name = domain.split('.')[0].title()
            analysis.append(f"• Suggested name: {suggested_name}")
            
        except Exception as e:
            analysis.append(f"• Error analyzing URL: {str(e)}")
        
        return '\n'.join(analysis)
    
    def _handle_source_name(self, chat_id: int, text: str) -> None:
        """Handle name input during source addition."""
        name = text.strip() if text.strip() else None
        
        # Use suggested name if empty
        if not name:
            analysis = self.user_states[chat_id].get('analysis', '')
            for line in analysis.split('\n'):
                if 'Suggested name:' in line:
                    name = line.split('Suggested name:')[1].strip()
                    break
        
        if not name:
            name = "News Source"
        
        # Store name and move to category step
        self.user_states[chat_id]['name'] = name
        self.user_states[chat_id]['state'] = 'awaiting_source_category'
        
        message = f"✅ Name set to: {name}\n\nPlease enter a category (e.g., Technology, Business, Sports, News)\n\nOr type 'skip' to skip:"
        self.send_message(chat_id, message)
    
    def _handle_source_category(self, chat_id: int, text: str) -> None:
        """Handle category input during source addition."""
        category = text.strip() if text.strip().lower() != 'skip' else None
        
        # Get stored data
        user_id = self.user_states[chat_id]['user_id']
        url = self.user_states[chat_id]['url']
        name = self.user_states[chat_id]['name']
        
        # Add source
        try:
            source_id = self.db.add_source(
                name=name,
                url=url,
                source_type='website',
                category=category,
                user_id=user_id
            )
            
            # Clear state
            del self.user_states[chat_id]
            
            message = f"✅ Source '{name}' added successfully!\n\nURL: {url}\nCategory: {category or 'N/A'}\n\nYou can view all your sources with /mysources"
            self.send_message(chat_id, message)
        except Exception as e:
            logger.error(f"Error adding source: {e}")
            self.send_message(chat_id, "❌ Failed to add source. Please try again.")
            del self.user_states[chat_id]
    
    def my_sources_command(self, chat_id: int) -> None:
        """Handle /mysources command."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register your@email.com")
            return
        
        sources = self._get_user_sources(user['id'])
        
        if not sources:
            self.send_message(chat_id, "You don't have any news sources yet.\n\nAdd one with /addsource")
            return
        
        message = "📰 Your News Sources\n\n"
        for i, source in enumerate(sources, 1):
            status = "✅" if source['enabled'] else "❌"
            message += f"{status} {i}. {source['name']}\n"
            message += f"   URL: {source['url']}\n"
            if source['rss_url']:
                message += f"   RSS: {source['rss_url']}\n"
            message += f"   Category: {source['category'] or 'N/A'}\n\n"
        
        message += f"Total: {len(sources)} sources\n\nUse /removesource to remove a source"
        
        self.send_message(chat_id, message)
    
    def remove_source_command(self, chat_id: int, args: list) -> None:
        """Handle /removesource command."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register your@email.com")
            return
        
        if not args or len(args) < 1:
            self.send_message(chat_id, "Usage: /removesource source_id\n\nUse /mysources to see your sources and their IDs")
            return
        
        try:
            source_id = int(args[0])
        except ValueError:
            self.send_message(chat_id, "❌ Invalid source ID. Use /mysources to see your sources")
            return
        
        # Check if source belongs to user
        source = self.db.get_source(source_id)
        if not source or source.get('user_id') != user['id']:
            self.send_message(chat_id, "❌ Source not found or doesn't belong to you")
            return
        
        # Remove source
        self._remove_source(source_id)
        self.send_message(chat_id, f"✅ Source '{source['name']}' removed successfully")
    
    def set_interval_command(self, chat_id: int, args: list) -> None:
        """Handle /setinterval command - interactive mode."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register")
            return
        
        # Start interactive interval setting
        self.user_states[chat_id] = {'state': 'awaiting_interval', 'user_id': user['id']}
        self.send_message(chat_id, "⏰ Please enter the check interval in hours (1-24):\n\nExample: 3")
    
    def _handle_interval(self, chat_id: int, text: str) -> None:
        """Handle interval input."""
        try:
            hours = int(text.strip())
            if hours < 1 or hours > 24:
                self.send_message(chat_id, "❌ Interval must be between 1 and 24 hours. Please try again:")
                return
        except ValueError:
            self.send_message(chat_id, "❌ Invalid interval. Please enter a number between 1 and 24:")
            return
        
        # Clear state
        del self.user_states[chat_id]
        
        self.send_message(chat_id, f"✅ Check interval set to {hours} hours\n\nNote: This will take effect on the next scheduler run")
    
    def cancel_command(self, chat_id: int) -> None:
        """Handle /cancel command - cancel current operation."""
        if chat_id in self.user_states:
            del self.user_states[chat_id]
            self.send_message(chat_id, "✅ Operation cancelled.")
        else:
            self.send_message(chat_id, "No active operation to cancel.")
    
    def stats_command(self, chat_id: int) -> None:
        """Handle /stats command."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register")
            return
        
        sources = self._get_user_sources(user['id'])
        articles = self._get_user_articles(user['id'])
        
        message = f"""📊 Your Statistics

👤 User: {user['name']}
📰 Sources: {len(sources)}
📄 Total Articles: {len(articles)}
� Last Check: {datetime.now().strftime('%Y-%m-%d %H:%M')}
"""
        
        recent_articles = sorted(articles, key=lambda x: x['created_at'], reverse=True)[:10]
        for article in recent_articles:
            message += f"\n• {article['title'][:50]}..."
        
        self.send_message(chat_id, message)
    
    def check_now_command(self, chat_id: int) -> None:
        """Handle /checknow command."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register")
            return
        
        self.send_message(chat_id, "🔍 Checking for new articles...\n\nThis may take a moment...")
        
        sources = self._get_user_sources(user['id'])
        all_new_articles = []
        
        for source in sources:
            try:
                articles = self._check_source(source, user['id'])
                all_new_articles.extend(articles)
            except Exception as e:
                logger.error(f"Error checking source {source['name']}: {e}")
        
        # Send articles to Telegram
        for article in all_new_articles:
            try:
                article['source_name'] = source['name']
                message = self._format_article_message(article)
                self.send_message(chat_id, message)
                time.sleep(1)  # Avoid rate limiting
            except Exception as e:
                logger.error(f"Error sending article: {e}")
        
        if all_new_articles:
            self.send_message(chat_id, f"✅ Found {len(all_new_articles)} new articles!")
        else:
            self.send_message(chat_id, "✅ No new articles found.")
    
    def _format_article_message(self, article: dict) -> str:
        """Format an article as a Telegram message."""
        from datetime import datetime
        
        message = f"*New Article*\n\n"
        message += f"📍 Source: {article.get('source_name', 'Unknown')}\n"
        message += f"📝 Title: {article.get('title', 'No title')}\n"
        
        # Add publication date/time
        pub_date = article.get('publication_date')
        if pub_date:
            try:
                if isinstance(pub_date, str):
                    dt = datetime.fromisoformat(pub_date.replace('Z', '+00:00'))
                else:
                    dt = pub_date
                message += f"🕐 Published: {dt.strftime('%Y-%m-%d %H:%M')}\n"
            except:
                message += f"🕐 Published: {pub_date}\n"
        else:
            # Use created_at if no publication date
            created_at = article.get('created_at')
            if created_at:
                try:
                    if isinstance(created_at, str):
                        dt = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
                    else:
                        dt = created_at
                    message += f"🕐 Added: {dt.strftime('%Y-%m-%d %H:%M')}\n"
                except:
                    pass
        
        if article.get('summary'):
            message += f"📄 Summary: {article['summary']}\n\n"
        else:
            message += "\n"
        
        message += f"🔗 [Read more]({article.get('url', '')})"
        
        return message
    
    def send_article_with_media(self, chat_id: int, article: dict) -> bool:
        """Send article with media (image/video) if available."""
        featured_image = article.get('featured_image')
        
        if featured_image:
            try:
                # Try to send as photo
                data = {
                    'chat_id': chat_id,
                    'photo': featured_image,
                    'caption': self._format_article_message(article),
                    'parse_mode': 'Markdown'
                }
                response = requests.post(f"{self.api_url}/sendPhoto", json=data, timeout=10)
                response.raise_for_status()
                result = response.json()
                logger.info(f"Sent article with photo: {result.get('ok', False)}")
                return result.get('ok', False)
            except Exception as e:
                logger.error(f"Error sending photo: {e}")
                # Fallback to text message
                return self.send_message(chat_id, self._format_article_message(article), parse_mode='Markdown')
        else:
            # No image, send as text
            return self.send_message(chat_id, self._format_article_message(article), parse_mode='Markdown')
    
    def check_today_command(self, chat_id: int) -> None:
        """Handle /checktoday command - check all articles from today."""
        self._check_articles_by_time(chat_id, hours=0, label="today")
    
    def check_yesterday_command(self, chat_id: int) -> None:
        """Handle /checkyesterday command - check articles from yesterday."""
        self._check_articles_by_time(chat_id, hours=24, label="yesterday")
    
    def check_12h_command(self, chat_id: int) -> None:
        """Handle /check12h command - check articles from last 12 hours."""
        self._check_articles_by_time(chat_id, hours=12, label="last 12 hours")
    
    def check_6h_command(self, chat_id: int) -> None:
        """Handle /check6h command - check articles from last 6 hours."""
        self._check_articles_by_time(chat_id, hours=6, label="last 6 hours")
    
    def check_1h_command(self, chat_id: int) -> None:
        """Handle /check1h command - check articles from last 1 hour."""
        self._check_articles_by_time(chat_id, hours=1, label="last 1 hour")
    
    def _check_articles_by_time(self, chat_id: int, hours: int, label: str) -> None:
        """Check articles from a specific time period."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register")
            return
        
        from datetime import datetime, timedelta
        import sqlite3
        
        # Calculate time threshold
        if hours == 0:  # Today
            threshold = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
        else:  # Last X hours
            threshold = datetime.now() - timedelta(hours=hours)
        
        threshold_str = threshold.strftime('%Y-%m-%d %H:%M:%S')
        
        self.send_message(chat_id, f"🔍 Checking articles from {label}...")
        
        # Get articles from database (limit to 10)
        with self.db._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT a.*, s.name as source_name 
                FROM articles a 
                JOIN sources s ON a.source_id = s.id 
                WHERE a.user_id = ? AND a.created_at >= ?
                ORDER BY a.created_at DESC
                LIMIT 10
            """, (user['id'], threshold_str))
            
            articles = []
            for row in cursor.fetchall():
                article = dict(row)
                articles.append(article)
        
        if not articles:
            self.send_message(chat_id, f"✅ No articles found from {label}.")
            return
        
        # Send articles to Telegram
        for article in articles:
            try:
                message = self._format_article_message(article)
                self.send_message(chat_id, message)
                time.sleep(1)  # Avoid rate limiting
            except Exception as e:
                logger.error(f"Error sending article: {e}")
        
        self.send_message(chat_id, f"✅ Found {len(articles)} articles from {label}.")
    
    # Keyword filtering commands
    def add_keyword_command(self, chat_id: int, args: list) -> None:
        """Handle /addkeyword command."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register")
            return
        
        if not args:
            self.send_message(chat_id, "❌ Please provide a keyword. Usage: /addkeyword bitcoin")
            return
        
        keyword = ' '.join(args).lower()
        try:
            self.db.add_keyword(user['id'], keyword)
            self.send_message(chat_id, f"✅ Keyword '{keyword}' added successfully!")
        except Exception as e:
            logger.error(f"Error adding keyword: {e}")
            self.send_message(chat_id, "❌ Failed to add keyword. It may already exist.")
    
    def remove_keyword_command(self, chat_id: int, args: list) -> None:
        """Handle /removekeyword command."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register")
            return
        
        if not args:
            self.send_message(chat_id, "❌ Please provide a keyword. Usage: /removekeyword bitcoin")
            return
        
        keyword = ' '.join(args).lower()
        if self.db.remove_keyword(user['id'], keyword):
            self.send_message(chat_id, f"✅ Keyword '{keyword}' removed successfully!")
        else:
            self.send_message(chat_id, f"❌ Keyword '{keyword}' not found.")
    
    def my_keywords_command(self, chat_id: int) -> None:
        """Handle /mykeywords command."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register")
            return
        
        keywords = self.db.get_user_keywords(user['id'])
        if not keywords:
            self.send_message(chat_id, "📝 You have no keywords set.\n\nUse /addkeyword to add keywords for filtering.")
        else:
            message = "📝 Your Keywords:\n\n"
            for keyword in keywords:
                message += f"• {keyword}\n"
            self.send_message(chat_id, message)
    
    # Bookmark commands
    def bookmark_command(self, chat_id: int, args: list) -> None:
        """Handle /bookmark command."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register")
            return
        
        if not args:
            self.send_message(chat_id, "❌ Please provide article ID. Usage: /bookmark article_id")
            return
        
        try:
            article_id = int(args[0])
            self.db.add_bookmark(user['id'], article_id)
            self.db.update_user_analytics(user['id'], articles_bookmarked=1)
            self.send_message(chat_id, f"✅ Article bookmarked successfully!")
        except ValueError:
            self.send_message(chat_id, "❌ Invalid article ID. Please provide a number.")
        except Exception as e:
            logger.error(f"Error bookmarking article: {e}")
            self.send_message(chat_id, "❌ Failed to bookmark article.")
    
    def unbookmark_command(self, chat_id: int, args: list) -> None:
        """Handle /unbookmark command."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register")
            return
        
        if not args:
            self.send_message(chat_id, "❌ Please provide article ID. Usage: /unbookmark article_id")
            return
        
        try:
            article_id = int(args[0])
            if self.db.remove_bookmark(user['id'], article_id):
                self.send_message(chat_id, f"✅ Bookmark removed successfully!")
            else:
                self.send_message(chat_id, f"❌ Bookmark not found.")
        except ValueError:
            self.send_message(chat_id, "❌ Invalid article ID. Please provide a number.")
        except Exception as e:
            logger.error(f"Error removing bookmark: {e}")
            self.send_message(chat_id, "❌ Failed to remove bookmark.")
    
    def my_bookmarks_command(self, chat_id: int) -> None:
        """Handle /mybookmarks command."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register")
            return
        
        bookmarks = self.db.get_user_bookmarks(user['id'])
        if not bookmarks:
            self.send_message(chat_id, "📚 You have no bookmarks.\n\nUse /bookmark article_id to save articles.")
        else:
            self.send_message(chat_id, f"📚 You have {len(bookmarks)} bookmarks:")
            for bookmark in bookmarks[:10]:  # Limit to 10
                message = self._format_article_message(bookmark)
                self.send_message(chat_id, message)
                time.sleep(1)
    
    # Search commands
    def search_command(self, chat_id: int, args: list) -> None:
        """Handle /search command."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register")
            return
        
        if not args:
            self.send_message(chat_id, "❌ Please provide search query. Usage: /search bitcoin")
            return
        
        query = ' '.join(args).lower()
        self.db.update_user_analytics(user['id'], searches_performed=1)
        
        # Search in article titles and summaries
        with self.db._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT a.*, s.name as source_name 
                FROM articles a 
                JOIN sources s ON a.source_id = s.id 
                WHERE a.user_id = ? 
                AND (LOWER(a.title) LIKE ? OR LOWER(a.summary) LIKE ?)
                ORDER BY a.created_at DESC
                LIMIT 10
            """, (user['id'], f'%{query}%', f'%{query}%'))
            
            articles = [dict(row) for row in cursor.fetchall()]
        
        if not articles:
            self.send_message(chat_id, f"❌ No articles found matching '{query}'.")
        else:
            self.send_message(chat_id, f"🔍 Found {len(articles)} articles matching '{query}':")
            for article in articles:
                message = self._format_article_message(article)
                self.send_message(chat_id, message)
                time.sleep(1)
    
    def search_from_command(self, chat_id: int, args: list) -> None:
        """Handle /searchfrom command."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register")
            return
        
        if len(args) < 2:
            self.send_message(chat_id, "❌ Please provide source name and query. Usage: /searchfrom source_name query")
            return
        
        source_name = args[0].lower()
        query = ' '.join(args[1:]).lower()
        self.db.update_user_analytics(user['id'], searches_performed=1)
        
        # Search in specific source
        with self.db._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT a.*, s.name as source_name 
                FROM articles a 
                JOIN sources s ON a.source_id = s.id 
                WHERE a.user_id = ? 
                AND LOWER(s.name) LIKE ?
                AND (LOWER(a.title) LIKE ? OR LOWER(a.summary) LIKE ?)
                ORDER BY a.created_at DESC
                LIMIT 10
            """, (user['id'], f'%{source_name}%', f'%{query}%', f'%{query}%'))
            
            articles = [dict(row) for row in cursor.fetchall()]
        
        if not articles:
            self.send_message(chat_id, f"❌ No articles found in '{source_name}' matching '{query}'.")
        else:
            self.send_message(chat_id, f"🔍 Found {len(articles)} articles in '{source_name}' matching '{query}':")
            for article in articles:
                message = self._format_article_message(article)
                self.send_message(chat_id, message)
                time.sleep(1)
    
    def _scheduled_check(self) -> None:
        """Scheduled check function called by the scheduler."""
        logger.info("Running scheduled check...")
        
        # Get all users
        with self.db._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT id, telegram_chat_id FROM users")
            users = cursor.fetchall()
        
        logger.info(f"Checking sources for {len(users)} users")
        
        for user_id, chat_id in users:
            try:
                # Get user's sources
                sources = self.db.get_user_sources(user_id)
                
                if not sources:
                    logger.info(f"User {user_id} has no sources, skipping")
                    continue
                
                logger.info(f"Checking {len(sources)} sources for user {user_id}")
                
                # Check each source for new articles
                all_new_articles = []
                for source in sources:
                    new_articles = self._check_source(source, user_id)
                    all_new_articles.extend(new_articles)
                
                # Send new articles to user
                if all_new_articles:
                    self.send_message(chat_id, f"🔔 You have {len(all_new_articles)} new articles!")
                    for article in all_new_articles:
                        message = self._format_article_message(article)
                        self.send_message(chat_id, message)
                        time.sleep(1)
                else:
                    # ALWAYS send articles - never leave user with nothing
                    from datetime import datetime, timedelta
                    
                    # Try today's articles first
                    threshold = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
                    
                    with self.db._get_connection() as conn:
                        cursor = conn.cursor()
                        cursor.execute("""
                            SELECT a.*, s.name as source_name 
                            FROM articles a 
                            JOIN sources s ON a.source_id = s.id 
                            WHERE a.user_id = ? AND a.created_at >= ?
                            ORDER BY a.created_at DESC
                            LIMIT 5
                        """, (user_id, threshold.strftime('%Y-%m-%d %H:%M:%S')))
                        
                        recent_articles = [dict(row) for row in cursor.fetchall()]
                    
                    if recent_articles:
                        self.send_message(chat_id, f"📰 Here are today's top articles ({len(recent_articles)}):")
                        for article in recent_articles:
                            message = self._format_article_message(article)
                            self.send_message(chat_id, message)
                            time.sleep(1)
                    else:
                        # If still no articles, send latest articles overall (guaranteed to have articles)
                        with self.db._get_connection() as conn:
                            cursor = conn.cursor()
                            cursor.execute("""
                                SELECT a.*, s.name as source_name 
                                FROM articles a 
                                JOIN sources s ON a.source_id = s.id 
                                WHERE a.user_id = ?
                                ORDER BY a.created_at DESC
                                LIMIT 5
                            """, (user_id,))
                            
                            latest_articles = [dict(row) for row in cursor.fetchall()]
                        
                        if latest_articles:
                            self.send_message(chat_id, f"📰 Here are the latest articles ({len(latest_articles)}):")
                            for article in latest_articles:
                                message = self._format_article_message(article)
                                self.send_message(chat_id, message)
                                time.sleep(1)
                        else:
                            # Fallback: fetch fresh articles immediately
                            self.send_message(chat_id, "🔄 Fetching fresh articles for you...")
                            for source in sources:
                                new_articles = self._check_source(source, user_id)
                                if new_articles:
                                    self.send_message(chat_id, f"🔔 Found {len(new_articles)} new articles!")
                                    for article in new_articles:
                                        message = self._format_article_message(article)
                                        self.send_message(chat_id, message)
                                        time.sleep(1)
                                    break
                
            except Exception as e:
                logger.error(f"Error in scheduled check for user {user_id}: {e}")
        
        logger.info("Scheduled check completed")
    
    # Scheduler control commands
    def start_scheduler_command(self, chat_id: int) -> None:
        """Handle /startscheduler command."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register")
            return
        
        try:
            self.scheduler.start()
            self.send_message(chat_id, "✅ Scheduler started successfully!")
        except Exception as e:
            logger.error(f"Error starting scheduler: {e}")
            self.send_message(chat_id, "❌ Failed to start scheduler. It may already be running.")
    
    def stop_scheduler_command(self, chat_id: int) -> None:
        """Handle /stopscheduler command."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register")
            return
        
        try:
            self.scheduler.stop()
            self.send_message(chat_id, "✅ Scheduler stopped successfully!")
        except Exception as e:
            logger.error(f"Error stopping scheduler: {e}")
            self.send_message(chat_id, "❌ Failed to stop scheduler.")
    
    def scheduler_status_command(self, chat_id: int) -> None:
        """Handle /schedulerstatus command."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register")
            return
        
        job_info = self.scheduler.get_job_info()
        if job_info:
            message = f"""📅 Scheduler Status

Status: {'Running' if self.scheduler.is_running else 'Stopped'}
Interval: {self.scheduler.check_interval_hours} hours
Next Run: {job_info.get('next_run_time', 'N/A')}"""
        else:
            message = "📅 Scheduler Status: Not configured"
        
        self.send_message(chat_id, message)
    
    # Category commands
    def category_command(self, chat_id: int, args: list) -> None:
        """Handle /category command - view articles by category."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register")
            return
        
        if not args:
            # Show available categories
            with self.db._get_connection() as conn:
                cursor = conn.cursor()
                cursor.execute("""
                    SELECT DISTINCT category FROM sources 
                    WHERE user_id = ? AND category IS NOT NULL
                """, (user['id'],))
                categories = [row[0] for row in cursor.fetchall()]
            
            if not categories:
                self.send_message(chat_id, "📂 No categories found. Use /setcategory to categorize your sources.")
            else:
                message = "📂 Available Categories:\n\n"
                for cat in categories:
                    message += f"• {cat}\n"
                message += "\nUsage: /category tech"
                self.send_message(chat_id, message)
        else:
            # Show articles from specific category
            category = ' '.join(args)
            with self.db._get_connection() as conn:
                cursor = conn.cursor()
                cursor.execute("""
                    SELECT a.*, s.name as source_name 
                    FROM articles a 
                    JOIN sources s ON a.source_id = s.id 
                    WHERE a.user_id = ? AND s.category = ?
                    ORDER BY a.created_at DESC
                    LIMIT 10
                """, (user['id'], category))
                
                articles = [dict(row) for row in cursor.fetchall()]
            
            if not articles:
                self.send_message(chat_id, f"❌ No articles found in category '{category}'.")
            else:
                self.send_message(chat_id, f"📂 Articles in '{category}':")
                for article in articles:
                    message = self._format_article_message(article)
                    self.send_message(chat_id, message)
                    time.sleep(1)
    
    def set_category_command(self, chat_id: int, args: list) -> None:
        """Handle /setcategory command - set category for a source."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register")
            return
        
        if len(args) < 2:
            self.send_message(chat_id, "❌ Usage: /setcategory source_id category")
            return
        
        try:
            source_id = int(args[0])
            category = ' '.join(args[1:])
            
            with self.db._get_connection() as conn:
                cursor = conn.cursor()
                cursor.execute("""
                    UPDATE sources 
                    SET category = ? 
                    WHERE id = ? AND user_id = ?
                """, (category, source_id, user['id']))
                
                if cursor.rowcount > 0:
                    self.send_message(chat_id, f"✅ Category set to '{category}' for source {source_id}.")
                else:
                    self.send_message(chat_id, "❌ Source not found or you don't have permission.")
        except ValueError:
            self.send_message(chat_id, "❌ Invalid source ID. Please provide a number.")
    
    # Source performance commands
    def source_stats_command(self, chat_id: int) -> None:
        """Handle /sourcestats command - view source performance metrics."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register")
            return
        
        performance_data = self.db.get_all_source_performance(user['id'])
        
        if not performance_data:
            self.send_message(chat_id, "📊 No performance data available yet. Check sources first.")
        else:
            message = "📊 Source Performance Metrics:\n\n"
            for perf in performance_data:
                success_rate = (perf['successful_checks'] / perf['total_checks'] * 100) if perf['total_checks'] > 0 else 0
                message += f"📰 {perf['source_name']}\n"
                message += f"   Total Checks: {perf['total_checks']}\n"
                message += f"   Success Rate: {success_rate:.1f}%\n"
                message += f"   Articles Found: {perf['total_articles_found']}\n"
                message += f"   Avg Response: {perf['avg_response_time']:.2f}s\n\n"
            
            self.send_message(chat_id, message)
    
    def reliable_sources_command(self, chat_id: int) -> None:
        """Handle /reliablesources command - show most reliable sources."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register")
            return
        
        performance_data = self.db.get_all_source_performance(user['id'])
        
        if not performance_data:
            self.send_message(chat_id, "📊 No performance data available yet.")
        else:
            # Sort by success rate
            reliable = sorted(performance_data, key=lambda x: (x['successful_checks'] / x['total_checks'] if x['total_checks'] > 0 else 0), reverse=True)
            
            message = "🏆 Most Reliable Sources:\n\n"
            for i, perf in enumerate(reliable[:5], 1):
                success_rate = (perf['successful_checks'] / perf['total_checks'] * 100) if perf['total_checks'] > 0 else 0
                message += f"{i}. {perf['source_name']} - {success_rate:.1f}% success\n"
            
            self.send_message(chat_id, message)
    
    # User analytics command
    def my_analytics_command(self, chat_id: int) -> None:
        """Handle /myanalytics command - view user analytics dashboard."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register")
            return
        
        analytics = self.db.get_user_analytics(user['id'])
        
        if not analytics:
            self.send_message(chat_id, "📊 No analytics data available yet. Start using the bot to track your activity!")
        else:
            message = f"""📊 Your Analytics Dashboard

👤 User: {user['name']}
📖 Articles Viewed: {analytics['articles_viewed']}
📚 Articles Bookmarked: {analytics['articles_bookmarked']}
🔍 Searches Performed: {analytics['searches_performed']}
🕐 Last Active: {analytics['last_active'] or 'Never'}"""
            
            self.send_message(chat_id, message)
    
    # Export commands
    def export_today_command(self, chat_id: int, args: list) -> None:
        """Handle /exporttoday command - export today's articles."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register")
            return
        
        format_type = args[0].lower() if args else 'json'
        
        if format_type not in ['json', 'csv']:
            self.send_message(chat_id, "❌ Invalid format. Use: /exporttoday json or /exporttoday csv")
            return
        
        from datetime import datetime, timedelta
        threshold = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
        
        with self.db._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT a.*, s.name as source_name 
                FROM articles a 
                JOIN sources s ON a.source_id = s.id 
                WHERE a.user_id = ? AND a.created_at >= ?
                ORDER BY a.created_at DESC
            """, (user['id'], threshold.strftime('%Y-%m-%d %H:%M:%S')))
            
            articles = [dict(row) for row in cursor.fetchall()]
        
        if not articles:
            self.send_message(chat_id, "❌ No articles found from today to export.")
            return
        
        # Generate export file
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"exports/articles_{user['name']}_{timestamp}.{format_type}"
        
        os.makedirs('exports', exist_ok=True)
        
        if format_type == 'json':
            import json
            with open(filename, 'w', encoding='utf-8') as f:
                json.dump(articles, f, indent=2, default=str)
        else:  # csv
            import csv
            with open(filename, 'w', newline='', encoding='utf-8') as f:
                writer = csv.DictWriter(f, fieldnames=articles[0].keys())
                writer.writeheader()
                writer.writerows(articles)
        
        self.send_message(chat_id, f"✅ Exported {len(articles)} articles to {filename}")
    
    def export_last_7days_command(self, chat_id: int, args: list) -> None:
        """Handle /exportlast7days command - export last 7 days articles."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register")
            return
        
        format_type = args[0].lower() if args else 'json'
        
        if format_type not in ['json', 'csv']:
            self.send_message(chat_id, "❌ Invalid format. Use: /exportlast7days json or /exportlast7days csv")
            return
        
        from datetime import datetime, timedelta
        threshold = datetime.now() - timedelta(days=7)
        
        with self.db._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT a.*, s.name as source_name 
                FROM articles a 
                JOIN sources s ON a.source_id = s.id 
                WHERE a.user_id = ? AND a.created_at >= ?
                ORDER BY a.created_at DESC
            """, (user['id'], threshold.strftime('%Y-%m-%d %H:%M:%S')))
            
            articles = [dict(row) for row in cursor.fetchall()]
        
        if not articles:
            self.send_message(chat_id, "❌ No articles found from last 7 days to export.")
            return
        
        # Generate export file
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"exports/articles_{user['name']}_7days_{timestamp}.{format_type}"
        
        os.makedirs('exports', exist_ok=True)
        
        if format_type == 'json':
            import json
            with open(filename, 'w', encoding='utf-8') as f:
                json.dump(articles, f, indent=2, default=str)
        else:  # csv
            import csv
            with open(filename, 'w', newline='', encoding='utf-8') as f:
                writer = csv.DictWriter(f, fieldnames=articles[0].keys())
                writer.writeheader()
                writer.writerows(articles)
        
        self.send_message(chat_id, f"✅ Exported {len(articles)} articles to {filename}")
    
    # Trending command
    def trending_command(self, chat_id: int) -> None:
        """Handle /trending command - view trending topics."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register")
            return
        
        from datetime import datetime, timedelta
        threshold = datetime.now() - timedelta(days=1)
        
        # Get recent articles and extract keywords from titles
        with self.db._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT title FROM articles 
                WHERE user_id = ? AND created_at >= ?
            """, (user['id'], threshold.strftime('%Y-%m-%d %H:%M:%S')))
            
            titles = [row[0] for row in cursor.fetchall()]
        
        if not titles:
            self.send_message(chat_id, "📈 No recent articles to analyze for trends.")
            return
        
        # Simple keyword extraction (count common words)
        from collections import Counter
        import re
        
        words = []
        for title in titles:
            # Extract words, remove common stop words
            title_words = re.findall(r'\b[a-zA-Z]{3,}\b', title.lower())
            stop_words = {'the', 'and', 'for', 'are', 'but', 'not', 'you', 'all', 'can', 'had', 'her', 'was', 'one', 'our', 'out', 'has', 'have', 'been', 'will', 'with', 'this', 'that', 'from', 'they', 'would', 'there', 'their', 'what', 'about', 'which', 'when', 'make', 'like', 'into', 'year', 'your', 'just', 'over', 'also', 'new', 'more', 'says'}
            words.extend([w for w in title_words if w not in stop_words])
        
        word_counts = Counter(words).most_common(10)
        
        if not word_counts:
            self.send_message(chat_id, "📈 Not enough data to determine trends.")
        else:
            message = "📈 Trending Topics (Last 24 Hours):\n\n"
            for i, (word, count) in enumerate(word_counts, 1):
                message += f"{i}. {word.capitalize()} ({count} mentions)\n"
            
            self.send_message(chat_id, message)
    
    def summarize_command(self, chat_id: int, args: list) -> None:
        """Handle /summarize command - AI-powered article summarization."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register")
            return
        
        if not args:
            # Summarize today's articles
            from datetime import datetime
            threshold = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
            
            with self.db._get_connection() as conn:
                cursor = conn.cursor()
                cursor.execute("""
                    SELECT id FROM articles 
                    WHERE user_id = ? AND created_at >= ?
                    ORDER BY created_at DESC
                    LIMIT 5
                """, (user['id'], threshold.strftime('%Y-%m-%d %H:%M:%S')))
                
                article_ids = [row[0] for row in cursor.fetchall()]
            
            if not article_ids:
                self.send_message(chat_id, "❌ No articles from today to summarize. Try checking for news first.")
                return
            
            self.send_message(chat_id, "🤖 Generating AI summary of today's articles...")
            self._ai_summarize_articles(chat_id, article_ids)
        else:
            # Summarize specific articles by ID
            try:
                article_ids = [int(arg) for arg in args]
                self.send_message(chat_id, f"🤖 Generating AI summary of {len(article_ids)} articles...")
                self._ai_summarize_articles(chat_id, article_ids)
            except ValueError:
                self.send_message(chat_id, "❌ Invalid article IDs. Please provide numbers.")
    
    def _get_user_by_chat_id(self, chat_id: int) -> Optional[dict]:
        """Get user by Telegram chat ID."""
        with self.db._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT id, name, telegram_chat_id FROM users WHERE telegram_chat_id = ?", (str(chat_id),))
            row = cursor.fetchone()
            if row:
                return {'id': row[0], 'name': row[1], 'telegram_chat_id': row[2]}
            return None
    
    def _get_user_by_name(self, name: str) -> Optional[dict]:
        """Get user by name."""
        with self.db._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT id, name, telegram_chat_id FROM users WHERE name = ?", (name,))
            row = cursor.fetchone()
            if row:
                return {'id': row[0], 'name': row[1], 'telegram_chat_id': row[2]}
            return None
    
    def _get_user_sources(self, user_id: int) -> list:
        """Get all sources for a user."""
        with self.db._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM sources WHERE user_id = ?", (user_id,))
            return [dict(row) for row in cursor.fetchall()]
    
    def _get_user_articles(self, user_id: int) -> list:
        """Get all articles for a user."""
        with self.db._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM articles WHERE user_id = ?", (user_id,))
            return [dict(row) for row in cursor.fetchall()]
    
    def _remove_source(self, source_id: int) -> None:
        """Remove a source."""
        with self.db._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM sources WHERE id = ?", (source_id,))
    
    def _check_source(self, source: dict, user_id: int) -> list:
        """Check a single source for new articles with performance tracking."""
        import time
        start_time = time.time()
        articles = []
        success = False
        
        try:
            # Rate limiting: check if source was checked recently
            last_check = source.get('last_checked')
            if last_check:
                from datetime import datetime, timedelta
                last_check_time = datetime.fromisoformat(last_check.replace('Z', '+00:00'))
                if datetime.now() - last_check_time < timedelta(minutes=5):
                    logger.info(f"Source {source['name']} was checked recently, skipping")
                    return []
            
            if source.get('rss_url'):
                articles = self.rss_parser.parse_feed(source['rss_url'], source['name'])
            else:
                articles = self.web_scraper.scrape_articles(source['url'], source['name'])
            
            new_articles = []
            for article in articles:
                # Check for duplicates using content hash
                title = article.get('title', '') or ''
                summary = article.get('summary', '') or ''
                content_hash = self.db._generate_content_hash(title + summary)
                article_id = self.db.add_article(
                    source_id=source['id'],
                    title=article['title'],
                    url=article['url'],
                    summary=article.get('summary'),
                    publication_date=article.get('publication_date'),
                    featured_image=article.get('featured_image'),
                    content=content_hash,
                    user_id=user_id
                )
                if article_id:
                    new_articles.append(article)
            
            self.db.update_source_last_checked(source['id'])
            success = True
            
            # Update performance metrics
            response_time = time.time() - start_time
            self.db.update_source_performance(user_id, source['id'], success, len(new_articles), response_time)
            
            return new_articles
            
        except Exception as e:
            logger.error(f"Error checking source {source['name']}: {e}")
            
            # Update performance metrics with failure
            response_time = time.time() - start_time
            self.db.update_source_performance(user_id, source['id'], False, 0, response_time)
            
            return []
    
    def _extract_name_from_url(self, url: str) -> str:
        """Extract a name from URL."""
        try:
            from urllib.parse import urlparse
            domain = urlparse(url).netloc
            return domain.replace('www.', '').split('.')[0].title()
        except:
            return "News Source"
    
    def _create_user(self, name: str, chat_id: int) -> int:
        """Create a new user."""
        with self.db._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO users (name, telegram_chat_id)
                VALUES (?, ?)
            """, (name, str(chat_id)))
            return cursor.lastrowid
    
    # Interest catalog commands
    def interest_command(self, chat_id: int, args: list) -> None:
        """Handle /interest command to mark an article as interesting."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register your@email.com")
            return
        
        if not args:
            self.send_message(chat_id, "❌ Please provide article ID. Usage: /interest article_id")
            return
        
        try:
            article_id = int(args[0])
        except ValueError:
            self.send_message(chat_id, "❌ Invalid article ID. Please provide a number.")
            return
        
        # Check if article exists and belongs to user
        with self.db._get_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT a.*, s.name as source_name 
                FROM articles a 
                JOIN sources s ON a.source_id = s.id 
                WHERE a.id = ? AND a.user_id = ?
            """, (article_id, user['id']))
            article = cursor.fetchone()
        
        if not article:
            self.send_message(chat_id, "❌ Article not found or doesn't belong to you.")
            return
        
        # Generate AI summary if available
        ai_summary = None
        if os.getenv('OPENAI_API_KEY'):
            try:
                article_title = article[3]
                article_summary = article[4] or ''
                ai_summary = self._generate_ai_summary(article_title, article_summary)
            except Exception as e:
                logger.error(f"Error generating AI summary: {e}")
        
        # Add to interests
        success = self.db.add_interest(user['id'], article_id, article[4], ai_summary)
        
        if success:
            self.send_message(chat_id, f"✅ Article marked as interesting!\n\n📝 {article[3][:100]}...")
            if ai_summary:
                self.send_message(chat_id, f"🤖 AI Summary: {ai_summary}")
        else:
            self.send_message(chat_id, "⚠️ Article already marked as interesting.")
    
    def my_interests_command(self, chat_id: int, args: list) -> None:
        """Handle /myinterests command to view daily interests."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register your@email.com")
            return
        
        # Get date from args or use today
        from datetime import datetime
        date = args[0] if args else datetime.now().strftime('%Y-%m-%d')
        
        interests = self.db.get_user_interests(user['id'], date)
        
        if not interests:
            self.send_message(chat_id, f"📝 No interests found for {date}.")
            return
        
        message = f"📝 Your interests for {date} ({len(interests)}):\n\n"
        for i, interest in enumerate(interests, 1):
            message += f"{i}. {interest['title']}\n"
            message += f"   📍 Source: {interest['source_name']}\n"
            if interest['ai_summary']:
                message += f"   🤖 AI: {interest['ai_summary'][:100]}...\n"
            message += f"   🔗 {interest['url']}\n\n"
        
        message += f"\nUse /exportdaily to generate a Word document with these interests."
        self.send_message(chat_id, message)
    
    def uninterest_command(self, chat_id: int, args: list) -> None:
        """Handle /uninterest command to remove from interests."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register your@email.com")
            return
        
        if not args:
            self.send_message(chat_id, "❌ Please provide article ID. Usage: /uninterest article_id")
            return
        
        try:
            article_id = int(args[0])
        except ValueError:
            self.send_message(chat_id, "❌ Invalid article ID. Please provide a number.")
            return
        
        success = self.db.remove_interest(user['id'], article_id)
        
        if success:
            self.send_message(chat_id, "✅ Article removed from interests.")
        else:
            self.send_message(chat_id, "❌ Article not found in your interests.")
    
    def export_daily_command(self, chat_id: int, args: list) -> None:
        """Handle /exportdaily command to generate Word document."""
        user = self._get_user_by_chat_id(chat_id)
        if not user:
            self.send_message(chat_id, "❌ Please register first with /register your@email.com")
            return
        
        from datetime import datetime
        date = args[0] if args else datetime.now().strftime('%Y-%m-%d')
        
        interests = self.db.get_user_interests(user['id'], date)
        
        if not interests:
            self.send_message(chat_id, f"📝 No interests found for {date}.")
            return
        
        # Generate Word document
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Title
            title = doc.add_heading(f'News Interests - {date}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add user name
            doc.add_paragraph(f"User: {user['name']}")
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph()
            
            # Add each interest
            for i, interest in enumerate(interests, 1):
                doc.add_heading(f"{i}. {interest['title']}", level=2)
                
                # Source
                p = doc.add_paragraph()
                p.add_run("Source: ").bold = True
                p.add_run(interest['source_name'])
                
                # URL
                p = doc.add_paragraph()
                p.add_run("URL: ").bold = True
                p.add_run(interest['url'])
                
                # Summary
                if interest['article_summary']:
                    p = doc.add_paragraph()
                    p.add_run("Summary: ").bold = True
                    p.add_run(interest['article_summary'])
                
                # AI Summary
                if interest['ai_summary']:
                    p = doc.add_paragraph()
                    p.add_run("AI Summary: ").bold = True
                    p.add_run(interest['ai_summary'])
                    p.run.font.color.rgb = RGBColor(0, 102, 204)
                
                doc.add_paragraph()
            
            # Save document
            filename = f"interests_{user['name']}_{date}.docx"
            filepath = os.path.join("exports", filename)
            os.makedirs("exports", exist_ok=True)
            doc.save(filepath)
            
            # Send file to user
            self.send_document(chat_id, filepath, filename)
            
            self.send_message(chat_id, f"✅ Word document generated: {filename}")
            
        except ImportError:
            self.send_message(chat_id, "❌ python-docx library not installed. Please install it with: pip install python-docx")
        except Exception as e:
            logger.error(f"Error generating Word document: {e}")
            self.send_message(chat_id, f"❌ Error generating document: {str(e)}")
    
    def _generate_ai_summary(self, title: str, summary: str) -> str:
        """Generate AI summary using OpenAI."""
        try:
            import openai
            openai.api_key = os.getenv('OPENAI_API_KEY')
            
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant that summarizes news articles concisely."},
                    {"role": "user", "content": f"Summarize this news article in 2-3 sentences:\n\nTitle: {title}\n\nSummary: {summary}"}
                ],
                max_tokens=150,
                temperature=0.7
            )
            
            return response.choices[0].message.content.strip()
        except Exception as e:
            logger.error(f"Error generating AI summary: {e}")
            return None
    
    def run(self) -> None:
        """Start the bot polling loop with HTTP health check server."""
        logger.info("Starting Telegram bot polling...")
        
        # Start the scheduler
        try:
            self.scheduler.start()
            logger.info("Scheduler started")
        except Exception as e:
            logger.error(f"Failed to start scheduler: {e}")
        
        # Start HTTP server for health checks in background thread
        port = int(os.environ.get('PORT', 10000))
        server = HTTPServer(('0.0.0.0', port), HealthCheckHandler)
        server_thread = threading.Thread(target=server.serve_forever, daemon=True)
        server_thread.start()
        logger.info(f"Health check server started on port {port}")
        
        while True:
            try:
                updates = self.get_updates()
                
                if updates:
                    logger.info(f"Received {len(updates)} updates")
                
                for update in updates:
                    if 'message' in update:
                        message = update['message']
                        chat_id = message['chat']['id']
                        text = message.get('text', '')
                        
                        logger.info(f"Processing message from {chat_id}: {text}")
                        
                        # Parse command
                        if text.startswith('/'):
                            parts = text.split()
                            command = parts[0][1:].lower()  # Remove / and lowercase
                            args = parts[1:] if len(parts) > 1 else []
                            logger.info(f"Command: {command}, Args: {args}")
                            self.handle_command(chat_id, command, args)
                        else:
                            # Handle non-command messages (conversation flow)
                            self.handle_message(chat_id, text)
                
                # Small delay to avoid excessive polling
                time.sleep(1)
                
            except KeyboardInterrupt:
                logger.info("Bot stopped by user")
                self.scheduler.stop()
                server.shutdown()
                break
            except Exception as e:
                logger.error(f"Error in polling loop: {e}")
                time.sleep(5)  # Wait before retrying


def main():
    """Main entry point for the bot."""
    import dotenv
    dotenv.load_dotenv()
    
    # Configure logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    bot_token = os.getenv('TELEGRAM_BOT_TOKEN')
    if not bot_token:
        logger.error("TELEGRAM_BOT_TOKEN not found in environment variables")
        return
    
    bot = NewsBot(bot_token)
    bot.run()


if __name__ == '__main__':
    main()
