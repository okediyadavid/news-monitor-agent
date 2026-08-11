import os
import sqlite3
import requests
from datetime import datetime
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv()

bot_token = os.getenv('TELEGRAM_BOT_TOKEN')
api_url = f"https://api.telegram.org/bot{bot_token}/sendMessage"
document_url = f"https://api.telegram.org/bot{bot_token}/sendDocument"

today = datetime.now().strftime('%Y-%m-%d')

# Users to test
users = [
    {'name': 'Dave', 'user_id': 1, 'chat_id': '1287986887'},
    {'name': 'Joshua', 'user_id': 3, 'chat_id': '2016570012'}
]

for user in users:
    user_id = user['user_id']
    chat_id = user['chat_id']
    name = user['name']
    
    print(f"\nTesting /exportdaily for {name}...")
    
    # Get user's interests for today
    conn = sqlite3.connect('data/news_monitor.db')
    cursor = conn.cursor()
    cursor.execute("""
        SELECT ui.*, a.title, a.url, a.summary as article_summary, s.name as source_name
        FROM user_interests ui
        JOIN articles a ON ui.article_id = a.id
        JOIN sources s ON a.source_id = s.id
        WHERE ui.user_id = ? AND DATE(ui.created_at) = ?
        ORDER BY ui.created_at DESC
    """, (user_id, today))
    
    interests = []
    for row in cursor.fetchall():
        interests.append({
            'id': row[0],
            'user_id': row[1],
            'article_id': row[2],
            'summary': row[3],
            'ai_summary': row[4],
            'created_at': row[5],
            'title': row[6],
            'url': row[7],
            'article_summary': row[8],
            'source_name': row[9]
        })
    conn.close()
    
    print(f"Found {len(interests)} interests for {name} on {today}")
    
    if not interests:
        print(f"No interests found for {name}")
        continue
    
    # Generate Word document
    doc = Document()
    
    # Title
    title = doc.add_heading(f'News Interests - {today}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add user name
    doc.add_paragraph(f"User: {name}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()
    
    # Add each interest
    for i, interest in enumerate(interests, 1):
        doc.add_heading(f"{i}. {interest['title']}", level=2)
        
        # Source
        p = doc.add_paragraph()
        p.add_run("Source: ").bold = True
        p.add_run(interest['source_name'])
        
        # URL
        p = doc.add_paragraph()
        p.add_run("URL: ").bold = True
        p.add_run(interest['url'])
        
        # Summary
        if interest['article_summary']:
            p = doc.add_paragraph()
            p.add_run("Summary: ").bold = True
            p.add_run(interest['article_summary'])
        
        # AI Summary
        if interest.get('ai_summary'):
            p = doc.add_paragraph()
            p.add_run("AI Summary: ").bold = True
            p.add_run(interest['ai_summary'])
            p.run.font.color.rgb = RGBColor(0, 102, 204)
        
        doc.add_paragraph()
    
    # Save document
    filename = f"interests_{name}_{today}.docx"
    filepath = os.path.join("exports", filename)
    os.makedirs("exports", exist_ok=True)
    doc.save(filepath)
    
    print(f"✅ Generated Word document: {filename}")
    
    # Send file to user
    try:
        with open(filepath, 'rb') as f:
            files = {'document': (filename, f)}
            data = {'chat_id': chat_id}
            
            response = requests.post(document_url, data=data, files=files, timeout=30)
            result = response.json()
            
            if result.get('ok'):
                print(f"✅ Word document sent to {name}")
            else:
                print(f"❌ Error sending to {name}: {result}")
    except Exception as e:
        print(f"❌ Error sending document to {name}: {e}")

print("\n✅ Export daily test completed for both users")
